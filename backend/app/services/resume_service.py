import pdfplumber
import docx
import io
import logging
from typing import Dict, List, Optional

logger = logging.getLogger(__name__)

nlp = None
try:
    import spacy
    try:
        nlp = spacy.load("en_core_web_sm")
    except OSError:
        logger.warning("spaCy model 'en_core_web_sm' not found. NLP extraction will be limited.")
        nlp = None
except Exception as e:
    logger.warning(f"spaCy unavailable: {e}. Resume NLP extraction will be disabled.")
    nlp = None

COMMON_SKILLS = [
    "python", "java", "c++", "c#", "javascript", "react", "node.js", "express",
    "fastapi", "django", "flask", "sql", "mysql", "postgresql", "mongodb",
    "docker", "kubernetes", "aws", "azure", "gcp", "machine learning",
    "deep learning", "nlp", "scikit-learn", "tensorflow", "pytorch",
    "html", "css", "git", "linux", "typescript", "angular", "vue"
]

def extract_text_from_pdf(file_bytes: bytes) -> str:
    """Takes PDF bytes and extracts text."""
    try:
        text = ""
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            for page in pdf.pages:
                selected_text = page.extract_text()
                if selected_text:
                    text += selected_text + "\n"
        return text
    except Exception as e:
        logger.error(f"Error extracting text from PDF: {e}")
        return ""

def extract_text_from_docx(file_bytes: bytes) -> str:
    """Takes DOCX bytes and extracts text from paragraphs and tables."""
    try:
        document = docx.Document(io.BytesIO(file_bytes))
        text_parts = []

        for paragraph in document.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)

        for table in document.tables:
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_text:
                    text_parts.append(" | ".join(row_text))

        return "\n".join(text_parts)
    except Exception as e:
        logger.error(f"Error extracting text from DOCX: {e}")
        return ""

def parse_resume(text: str) -> Dict:
    """Parses text to find skills, education, and experience."""
    logger.info("Parsing resume content...")
    skills_found = set()
    
    lower_text = text.lower()
    for skill in COMMON_SKILLS:
        if skill in lower_text:
            skills_found.add(skill.capitalize())
            
    education = []
    experience = []
    
    if nlp is not None:
        doc = nlp(text)
        for ent in doc.ents:
            if ent.label_ == "ORG":
                experience.append(ent.text)

    experience = list(set([exp.strip() for exp in experience if len(exp.strip()) > 3]))[:10]
    
    logger.info(f"Resume parsing complete. Found {len(skills_found)} skills.")
    return {
        "skills": list(skills_found),
        "education": education,
        "experience": experience
    }
